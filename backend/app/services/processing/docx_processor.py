"""
DOCX Processor
"""

import os
import logging
from typing import Dict, Any
from docx import Document

from app.services.processing.base import DocumentProcessor

logger = logging.getLogger(__name__)


class DocxProcessor(DocumentProcessor):
    """Process DOCX Word documents."""

    def __init__(self, file_path: str, filename: str):
        super().__init__(file_path, filename)

    def can_process(self) -> bool:
        return os.path.exists(self.file_path) and self.file_path.lower().endswith('.docx')

    def extract_text(self) -> str:
        doc = Document(self.file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(' | '.join(cells))
        return '\n'.join(paragraphs)

    def extract_metadata(self) -> Dict[str, Any]:
        doc = Document(self.file_path)
        core = doc.core_properties
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return {
            'title': core.title or '',
            'author': core.author or '',
            'subject': core.subject or '',
            'word_count': sum(len(p.split()) for p in paragraphs),
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables),
        }

    def extract_structure(self) -> Dict[str, Any]:
        doc = Document(self.file_path)
        headings = []
        for p in doc.paragraphs:
            if p.style and p.style.name.startswith('Heading'):
                headings.append({
                    'text': p.text.strip(),
                    'level': int(p.style.name.replace('Heading ', '') or '1'),
                })
        return {
            'headings': headings,
            'has_tables': len(doc.tables) > 0,
            'has_images': bool(doc.inline_shapes),
        }
